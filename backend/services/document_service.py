import subprocess
from pathlib import Path

from utils.file_utils import temp_path
from utils.platform_utils import get_libreoffice_path
from utils.zip_utils import bundle_to_zip


# ── LibreOffice helper ────────────────────────────────────────────────────────

def _libreoffice_convert(inp: Path, output_dir: Path, target_fmt: str) -> Path:
    """Convert a file to target_fmt via LibreOffice headless."""
    try:
        soffice = get_libreoffice_path()
    except RuntimeError as e:
        raise RuntimeError(str(e))

    result = subprocess.run(
        [soffice, "--headless", "--convert-to", target_fmt, "--outdir", str(output_dir), str(inp)],
        capture_output=True,
        text=True,
        timeout=120,
    )
    if result.returncode != 0:
        raise RuntimeError(f"LibreOffice error: {result.stderr[-600:]}")

    out = output_dir / (inp.stem + f".{target_fmt}")
    if not out.exists():
        # LibreOffice sometimes uses a different stem
        candidates = list(output_dir.glob(f"*.{target_fmt}"))
        if not candidates:
            raise RuntimeError(f"LibreOffice produced no .{target_fmt} output")
        out = candidates[0]
    return out


# ── PDF operations (PyMuPDF) ──────────────────────────────────────────────────

def pdf_to_word(inp: Path, out: Path) -> None:
    try:
        from pdf2docx import Converter  # type: ignore[import-untyped]
    except ImportError:
        raise RuntimeError("pdf2docx is required. Run: pip install pdf2docx")
    cv = Converter(str(inp))
    cv.convert(str(out))
    cv.close()


def pdf_to_txt(inp: Path, out: Path) -> None:
    import fitz  # PyMuPDF
    doc = fitz.open(str(inp))
    text = "\n".join(page.get_text() for page in doc)
    doc.close()
    out.write_text(text, encoding="utf-8")


def pdf_to_html(inp: Path, out: Path) -> None:
    import fitz
    doc = fitz.open(str(inp))
    html_parts = [
        '<!DOCTYPE html><html><head><meta charset="utf-8"></head><body>'
    ]
    for page in doc:
        html_parts.append(page.get_text("html"))
    html_parts.append("</body></html>")
    doc.close()
    out.write_text("\n".join(html_parts), encoding="utf-8")


def pdf_to_images_zip(inp: Path, out_zip: Path, dpi: int = 150) -> None:
    import fitz
    doc = fitz.open(str(inp))
    img_paths: list[tuple[Path, str]] = []
    mat = fitz.Matrix(dpi / 72, dpi / 72)
    for i, page in enumerate(doc):
        pix = page.get_pixmap(matrix=mat)
        p = temp_path("png")
        pix.save(str(p))
        img_paths.append((p, f"page_{i + 1:04d}.png"))
    doc.close()
    bundle_to_zip(img_paths, out_zip)
    for p, _ in img_paths:
        p.unlink(missing_ok=True)


def pdf_merge(input_paths: list[Path], out: Path) -> None:
    import fitz
    merged = fitz.open()
    for p in input_paths:
        src = fitz.open(str(p))
        merged.insert_pdf(src)
        src.close()
    merged.save(str(out))
    merged.close()


def pdf_split(inp: Path, out_zip: Path) -> None:
    import fitz
    doc = fitz.open(str(inp))
    parts: list[tuple[Path, str]] = []
    for i in range(len(doc)):
        p = temp_path("pdf")
        page_doc = fitz.open()
        page_doc.insert_pdf(doc, from_page=i, to_page=i)
        page_doc.save(str(p))
        page_doc.close()
        parts.append((p, f"page_{i + 1:04d}.pdf"))
    doc.close()
    bundle_to_zip(parts, out_zip)
    for p, _ in parts:
        p.unlink(missing_ok=True)


def pdf_compress(inp: Path, out: Path) -> None:
    import fitz
    doc = fitz.open(str(inp))
    doc.save(str(out), garbage=4, deflate=True, clean=True)
    doc.close()


# ── Office → PDF (LibreOffice) ────────────────────────────────────────────────

def office_to_pdf(inp: Path, out: Path) -> None:
    """Convert DOCX, ODT, RTF, PPT, ODP, XLSX, or .pages to PDF."""
    tmp_dir = inp.parent
    pdf_path = _libreoffice_convert(inp, tmp_dir, "pdf")
    if pdf_path != out:
        pdf_path.rename(out)


def office_convert(inp: Path, out: Path, target_fmt: str) -> None:
    """General LibreOffice conversion (docx→odt, odt→docx, etc.)."""
    tmp_dir = inp.parent
    result_path = _libreoffice_convert(inp, tmp_dir, target_fmt)
    if result_path != out:
        result_path.rename(out)


# ── HTML / Markdown ───────────────────────────────────────────────────────────

_MD_CSS = (
    "body{font-family:system-ui,sans-serif;max-width:800px;margin:40px auto;"
    "padding:0 20px;line-height:1.65;color:#111}"
    "h1,h2,h3{font-weight:600;margin-top:1.5em}"
    "code{background:#f5f5f5;padding:2px 5px;border-radius:3px;font-size:.875em}"
    "pre{background:#f5f5f5;padding:16px;overflow-x:auto}"
    "pre code{background:none;padding:0}"
    "table{border-collapse:collapse;width:100%}"
    "th,td{border:1px solid #ddd;padding:8px 12px;text-align:left}"
    "th{background:#f5f5f5}"
    "blockquote{border-left:3px solid #0066FF;margin:0;padding-left:16px;color:#555}"
)


def md_to_html_str(md_text: str) -> str:
    import markdown as md_lib  # type: ignore[import-untyped]
    body = md_lib.markdown(md_text, extensions=["tables", "fenced_code", "toc"])
    return f'<!DOCTYPE html><html><head><meta charset="utf-8"><style>{_MD_CSS}</style></head><body>{body}</body></html>'


def html_to_pdf(html_content: str, out: Path) -> None:
    try:
        from weasyprint import HTML  # type: ignore[import-untyped]
    except ImportError:
        raise RuntimeError("WeasyPrint is required. Run: pip install WeasyPrint")
    HTML(string=html_content).write_pdf(str(out))


def md_to_pdf(md_text: str, out: Path) -> None:
    html_to_pdf(md_to_html_str(md_text), out)


def md_to_docx(md_text: str, out: Path) -> None:
    """Convert Markdown to DOCX via HTML → LibreOffice."""
    html_content = md_to_html_str(md_text)
    tmp_html = temp_path("html")
    try:
        tmp_html.write_text(html_content, encoding="utf-8")
        office_convert(tmp_html, out, "docx")
    finally:
        tmp_html.unlink(missing_ok=True)


def txt_to_pdf(txt_path: Path, out: Path) -> None:
    text = txt_path.read_text(encoding="utf-8", errors="replace")
    escaped = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    html = (
        f'<!DOCTYPE html><html><head><meta charset="utf-8">'
        f'<style>body{{font-family:monospace;white-space:pre-wrap;padding:2rem}}</style>'
        f'</head><body>{escaped}</body></html>'
    )
    html_to_pdf(html, out)


# ── DOCX extraction (python-docx) ─────────────────────────────────────────────

def docx_to_txt(inp: Path, out: Path) -> None:
    try:
        from docx import Document  # type: ignore[import-untyped]
    except ImportError:
        raise RuntimeError("python-docx is required. Run: pip install python-docx")
    doc = Document(str(inp))
    text = "\n".join(para.text for para in doc.paragraphs)
    out.write_text(text, encoding="utf-8")


def docx_to_html(inp: Path, out: Path) -> None:
    try:
        from docx import Document  # type: ignore[import-untyped]
    except ImportError:
        raise RuntimeError("python-docx is required.")
    doc = Document(str(inp))
    parts = ['<!DOCTYPE html><html><head><meta charset="utf-8"></head><body>']
    for para in doc.paragraphs:
        style = para.style.name.lower()
        if style.startswith("heading"):
            lvl = style.replace("heading", "").strip() or "1"
            parts.append(f"<h{lvl}>{para.text}</h{lvl}>")
        elif para.text.strip():
            parts.append(f"<p>{para.text}</p>")
    parts.append("</body></html>")
    out.write_text("\n".join(parts), encoding="utf-8")
